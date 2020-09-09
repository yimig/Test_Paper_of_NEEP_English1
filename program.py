from docx import *
import json

with open("source.json", "r", encoding='utf8') as f:
    raw = f.read()
data = json.loads(raw)
file = Document("temp.docx")
section_num = 0
for exam in data['exams']:
    year = str(exam["year"])
    for topic in exam["read"]:
        content = topic["question"]
        num = str(topic["num"])
        file.add_paragraph(content)
        file.add_section()
        section_num += 1
        file.sections[section_num].header.paragraphs[0].text = year +"年 Text"+num
        file.sections[section_num].header.is_linked_to_previous = False
file.save("work\\"+"all.docx")
